import os
os.environ["OMP_NUM_THREADS"] = "1"
os.environ["MKL_NUM_THREADS"] = "1"
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
import gc
import requests
from requests_toolbelt.multipart.encoder import MultipartEncoder
import json
from pdf2image import convert_from_path, pdfinfo_from_path
from fastapi import FastAPI, File, UploadFile
from fastapi.responses import JSONResponse
from paddleocr import PPStructure
import pandas as pd
from PIL import Image
import tempfile
from docx2pdf import convert as docx2pdf_convert
import numpy as np
import paddle
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
from markdown2 import markdown
paddle.set_device("cpu")

app = FastAPI()

UPLOAD_DIR = "./uploaded_files"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Dify 接口参数
DIFY_API_URL = "https://dify.medicalunion.cn/v1/workflows/run"
DIFY_API_KEY = "app-PqDDjiDo6tX8VmleHzAGJSd8"

# 上传文件到 Dify 文件服务
def upload_file_to_dify(image_path):
    m = MultipartEncoder(
        fields={
            "file": ("screenshot.png", open(image_path, "rb"), "image/png"),
            "user": "ocr-user"
        }
    )

    headers = {
        "Authorization": f"Bearer {DIFY_API_KEY}",
        "Content-Type": m.content_type
    }

    response = requests.post("https://dify.medicalunion.cn/v1/files/upload", headers=headers, data=m)
    if response.status_code in (200, 201):
        file_id = response.json().get("id")
        print(f"[上传成功] 文件 ID: {file_id}")
        return file_id
    else:
        print(f"[UPLOAD ERROR] 状态码: {response.status_code}, 响应: {response.text}")
        return None

def get_image_description_dify(image_path):
    file_id = upload_file_to_dify(image_path)
    if not file_id:
        return "上传失败，未获取文件 ID"

    headers = {
        "Authorization": f"Bearer {DIFY_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "inputs": {
            "file": [
                {
                    "transfer_method": "local_file",
                    "upload_file_id": file_id,
                    "type": "image"
                }
            ]
        },
        "response_mode": "blocking",
        "user": "ocr-user"
    }

    response = requests.post(DIFY_API_URL, headers=headers, data=json.dumps(payload))

    if response.status_code != 200:
        print(f"[DIFY ERROR] 状态码: {response.status_code}, 响应: {response.text}")
        return "调用失败：" + response.text

    response_json = response.json()
    print("[DEBUG] 模型完整响应：", json.dumps(response_json, ensure_ascii=False, indent=2))
    text = ""
    try:
        text = response_json.get("data", {}).get("outputs", {}).get("text", "")
    except Exception as e:
        print(f"[解析错误] 无法提取 text 字段：{e}")
    text = text.strip() if isinstance(text, str) else ""
    return text if text else "未能识别"

def save_crop(image, bbox, save_path):
    if not isinstance(bbox, (list, tuple)):
        print(f"[WARNING] bbox 不是列表格式: {bbox}，跳过截图")
        return None
    if not bbox or len(bbox) < 4:
        print(f"[WARNING] 无效的 bbox: {bbox}，跳过截图")
        return None
    if isinstance(bbox[0], (list, tuple)):
        x_min = min([point[0] for point in bbox])
        y_min = min([point[1] for point in bbox])
        x_max = max([point[0] for point in bbox])
        y_max = max([point[1] for point in bbox])
    else:
        x_min, y_min, x_max, y_max = bbox

    cropped_image = image.crop((x_min, y_min, x_max, y_max))
    cropped_image.save(save_path)
    return save_path

@app.post("/process-structured/")
async def process_structured(file: UploadFile = File(...)):
    ocr = PPStructure(layout=True, show_log=True)

    base_name = os.path.splitext(file.filename)[0]
    file_dir = os.path.join(UPLOAD_DIR, base_name)
    os.makedirs(file_dir, exist_ok=True)

    file_location = os.path.join(file_dir, file.filename)
    with open(file_location, "wb") as f:
        f.write(await file.read())

    data = []

    if file.filename.lower().endswith(('.doc', '.docx')):
        temp_pdf_dir = tempfile.mkdtemp()
        pdf_output_path = os.path.join(temp_pdf_dir, f"{os.path.splitext(file.filename)[0]}.pdf")
        try:
            docx2pdf_convert(file_location, pdf_output_path)
            print(f"[INFO] Word 文件已转换为 PDF: {pdf_output_path}")
            num_pages = pdfinfo_from_path(pdf_output_path)["Pages"]
            for page_num in range(num_pages):
                print(f"[进度] 正在处理 Word->PDF 第 {page_num+1}/{num_pages} 页")
                page_image = convert_from_path(pdf_output_path, first_page=page_num+1, last_page=page_num+1)[0]
                image_np = np.array(page_image)
                result = ocr(image_np)
                result.sort(key=lambda x: (x.get("bbox", [0, 0, 0, 0])[1], x.get("bbox", [0, 0, 0, 0])[0]))
                for idx, item in enumerate(result):
                    block_type = item['type']
                    bbox = item['bbox']
                    print(f"[DEBUG] Word->PDF Page {page_num+1}, Block type: {block_type}, BBox: {bbox}")

                    if block_type == 'text':
                        text_content = ' '.join([res['text'] for res in item['res']])
                        data.append({
                            '页面': page_num + 1,
                            '类型': 'text',
                            '内容': text_content,
                            # '位置': bbox
                        })
                    elif block_type in ['table', 'figure']:
                        crop_path = os.path.join(file_dir, f"{block_type}_wordpage{page_num+1}_{idx}.png")
                        saved_path = save_crop(page_image, bbox, crop_path)
                        if saved_path:
                            description = get_image_description_dify(saved_path)
                            data.append({
                                '页面': page_num + 1,
                                '类型': block_type,
                                '截图路径': saved_path,
                                # '位置': bbox,
                                '识别描述': description
                            })
                del page_image
                del image_np
            del result
        finally:
            import shutil
            shutil.rmtree(temp_pdf_dir)
        del ocr

    elif file.filename.lower().endswith(".pdf"):
        num_pages = pdfinfo_from_path(file_location)["Pages"]
        for page_num in range(num_pages):
            print(f"[进度] 正在处理 PDF 第 {page_num+1}/{num_pages} 页")
            page_image = convert_from_path(file_location, first_page=page_num+1, last_page=page_num+1)[0]
            image_np = np.array(page_image)
            result = ocr(image_np)
            result.sort(key=lambda x: (x.get("bbox", [0, 0, 0, 0])[1], x.get("bbox", [0, 0, 0, 0])[0]))

            for idx, item in enumerate(result):
                block_type = item['type']
                bbox = item['bbox']
                print(f"[DEBUG] PDF Page {page_num+1}, Block type: {block_type}, BBox: {bbox}")

                if block_type == 'text':
                    text_content = ' '.join([res['text'] for res in item['res']])
                    data.append({
                        '页面': page_num + 1,
                        '类型': 'text',
                        '内容': text_content,
                        # '位置': bbox
                    })
                elif block_type in ['table', 'figure']:
                    crop_path = os.path.join(file_dir, f"{block_type}_page{page_num+1}_{idx}.png")
                    saved_path = save_crop(page_image, bbox, crop_path)
                    if saved_path:
                        description = get_image_description_dify(saved_path)
                        data.append({
                            '页面': page_num + 1,
                            '类型': block_type,
                            '截图路径': saved_path,
                            # '位置': bbox,
                            '识别描述': description
                        })
            del page_image
            del image_np
        del result
        del ocr

    else:
        print(f"[进度] 正在处理单张图片文件: {file.filename}")
        ocr = PPStructure(layout=True, show_log=True)
        image = Image.open(file_location)
        result = ocr(image)
        result.sort(key=lambda x: (x.get("bbox", [0, 0, 0, 0])[1], x.get("bbox", [0, 0, 0, 0])[0]))

        for idx, item in enumerate(result):
            block_type = item['type']
            bbox = item['bbox']
            print(f"[DEBUG] Block type: {block_type}, BBox: {bbox}")

            if block_type == 'text':
                text_content = ' '.join([res['text'] for res in item['res']])
                data.append({
                    '类型': 'text',
                    '内容': text_content,
                    # '位置': bbox
                })
            elif block_type in ['table', 'figure']:
                crop_path = os.path.join(file_dir, f"{block_type}_{idx}.png")
                saved_path = save_crop(image, bbox, crop_path)
                if saved_path:
                    description = get_image_description_dify(saved_path)
                    data.append({
                        '类型': block_type,
                        '截图路径': saved_path,
                        # '位置': bbox,
                        '识别描述': description
                    })
        del image
        del result
        del ocr

    for var in ['image', 'image_np', 'result', 'ocr']:
        if var in locals():
            del locals()[var]
    gc.collect()

    # ✅ 生成 Markdown 文件
    md_lines = []

    for item in data:
        page = item.get("页面", "?")
        block_type = item.get("类型", "")

        section = []

        # # 插入截图（如果有）
        # if "截图路径" in item:
        #     rel_path = os.path.relpath(item["截图路径"], start=os.path.join("./md_outputs", base_name))
        #     section.append(f"![截图]({rel_path})\n")

        # 插入识别描述（来自大模型）
        desc = item.get("识别描述", "").strip()
        if desc:
            # section.append("**图片描述：**\n")
            section.append(desc)

        # 插入 OCR 原始文字（仅 text 类型）
        if block_type == "text":
            content = item.get("内容", "").strip()
            if content:
                # section.append("**OCR识别文本：**\n")
                section.append(content)

        # 如果有内容才添加标题与分隔符
        if section:
            # md_lines.append(f"### 第{page}页 - {block_type.capitalize()}\n")
            md_lines.extend(section)
            md_lines.append("\n---\n")

    if md_lines:
        md_output_dir = os.path.join("./md_outputs", base_name)
        os.makedirs(md_output_dir, exist_ok=True)
        md_output_path = os.path.join(md_output_dir, f"{base_name}.md")
        with open(md_output_path, "w", encoding="utf-8") as f_md:
            f_md.write("\n".join(md_lines))
        print(f"[输出] Markdown 文件已生成: {md_output_path}")

        # 根据上传的文件类型生成对应格式的输出文档
        def convert_md_to_docx(md_path, output_path):
            with open(md_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()

            doc = Document()
            for line in lines:
                line = line.strip()
                if line.startswith("![](") and line.endswith(")"):
                    # 处理图片路径
                    image_path = line[4:-1]
                    if os.path.exists(image_path):
                        doc.add_picture(image_path, width=Inches(6))
                elif line.startswith("---"):
                    doc.add_paragraph().add_run(" ").add_break()
                else:
                    doc.add_paragraph(line)
            doc.save(output_path)

        def convert_md_to_pdf(md_path, output_path):
            pdf = FPDF()
            pdf.add_page()

            # 添加中文字体支持
            font_path = "/System/Library/Fonts/STHeiti Medium.ttc"  # macOS 默认中文字体
            pdf.add_font("heiti", "", font_path, uni=True)
            pdf.set_font("heiti", size=12)

            def safe_multicell(text):
                max_width = pdf.w - 2 * pdf.l_margin
                try:
                    pdf.multi_cell(0, 10, text)
                except Exception as e:
                    print(f"[PDF 渲染警告] 当前行过宽，逐字符强制换行：{text[:50]}...")
                    # 将文本中不能换行的长串切分成字符逐行输出
                    lines = []
                    current_line = ''
                    for char in text:
                        if pdf.get_string_width(current_line + char) < max_width:
                            current_line += char
                        else:
                            lines.append(current_line)
                            current_line = char
                    if current_line:
                        lines.append(current_line)

                    for l in lines:
                        try:
                            pdf.multi_cell(0, 10, l)
                        except Exception as e2:
                            print(f"[致命字符错误] 渲染失败：{repr(l)} - {e2}")
                            pdf.multi_cell(0, 10, '[Unrenderable Content]')

            with open(md_path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        pdf.ln(8)
                    else:
                        safe_multicell(line)

            pdf.output(output_path)

        if file.filename.lower().endswith((".doc", ".docx")):
            docx_output_path = os.path.join(md_output_dir, f"{base_name}_output.docx")
            convert_md_to_docx(md_output_path, docx_output_path)
            print(f"[输出] Word 文件已生成: {docx_output_path}")
        elif file.filename.lower().endswith(".pdf"):
            pdf_output_path = os.path.join(md_output_dir, f"{base_name}_output.pdf")
            convert_md_to_pdf(md_output_path, pdf_output_path)
            print(f"[输出] PDF 文件已生成: {pdf_output_path}")

    return JSONResponse(content={"results": data})

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("ocr_test:app", host="0.0.0.0", port=1408, reload=True)